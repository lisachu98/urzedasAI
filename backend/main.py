import os
import base64
import mimetypes
import logging
import asyncio
from pathlib import Path
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import Optional
import anthropic
import openai
from google import genai
from google.genai import types as gtypes
import fitz
import re
import time
from functools import lru_cache

# ── logging ───────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
log = logging.getLogger("docai")

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

DOCS_ROOT = os.environ.get("DOCS_ROOT", "/docs")
MAX_FILE_SIZE = 50 * 1024 * 1024    # 50 MB per file
MAX_TOTAL_SIZE = 200 * 1024 * 1024  # 200 MB total context

# ── singleton API clients (async) ────────────────────────────────────────────

@lru_cache(maxsize=1)
def get_anthropic_client() -> anthropic.AsyncAnthropic:
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not set")
    return anthropic.AsyncAnthropic(api_key=api_key)

@lru_cache(maxsize=1)
def get_openai_client() -> openai.AsyncOpenAI:
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY not set")
    return openai.AsyncOpenAI(api_key=api_key)

@lru_cache(maxsize=1)
def get_gemini_client() -> genai.Client:
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="GEMINI_API_KEY not set")
    return genai.Client(api_key=api_key)

@lru_cache(maxsize=1)
def get_xai_client() -> openai.AsyncOpenAI:
    api_key = os.environ.get("XAI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="XAI_API_KEY not set")
    return openai.AsyncOpenAI(api_key=api_key, base_url="https://api.x.ai/v1")

# ── model listing (cached) ───────────────────────────────────────────────────

_models_cache: dict = {"data": None, "ts": 0.0}
_CACHE_TTL = 3600  # 1 hour
_OAI_DATED_RE = re.compile(r"\d{4,}")  # filters out dated snapshots like gpt-4o-2024-08-06


async def fetch_available_models() -> dict:
    """Query each provider for available models. Results cached for 1 hour."""
    now = time.time()
    if _models_cache["data"] and (now - _models_cache["ts"]) < _CACHE_TTL:
        return _models_cache["data"]

    result: dict[str, list] = {}

    # ── Anthropic ──
    try:
        client = get_anthropic_client()
        page = await client.models.list(limit=100)
        result["anthropic"] = sorted(
            [{"id": m.id, "name": m.display_name} for m in page.data],
            key=lambda x: x["name"],
        )
        log.info("Loaded %d Anthropic models", len(result["anthropic"]))
    except Exception as e:
        log.warning("Failed to load Anthropic models: %s", e)
        result["anthropic"] = []

    # ── OpenAI ──
    try:
        client = get_openai_client()
        oai: list[dict] = []
        oai_page = await client.models.list()
        for m in oai_page.data:
            mid = m.id
            # only chat-capable model families
            if not (mid.startswith("gpt-") or re.match(r"^o\d", mid) or mid.startswith("chatgpt")):
                continue
            # skip non-chat variants
            if any(kw in mid for kw in ("realtime", "audio", "transcribe", "instruct", "search")):
                continue
            # skip dated snapshots (e.g. gpt-4o-2024-08-06)
            if _OAI_DATED_RE.search(mid):
                continue
            oai.append({"id": mid, "name": mid})
        result["openai"] = sorted(oai, key=lambda x: x["name"])
        log.info("Loaded %d OpenAI models", len(result["openai"]))
    except Exception as e:
        log.warning("Failed to load OpenAI models: %s", e)
        result["openai"] = []

    # ── Gemini (sync SDK — run in thread) ──
    try:
        client = get_gemini_client()
        def _list_gemini():
            gem: list[dict] = []
            for m in client.models.list():
                model_id = m.name.replace("models/", "") if m.name else ""
                if not model_id.startswith("gemini"):
                    continue
                gem.append({"id": model_id, "name": m.display_name or model_id})
            return sorted(gem, key=lambda x: x["name"])
        result["gemini"] = await asyncio.to_thread(_list_gemini)
        log.info("Loaded %d Gemini models", len(result["gemini"]))
    except Exception as e:
        log.warning("Failed to load Gemini models: %s", e)
        result["gemini"] = []

    # ── xAI ──
    try:
        client = get_xai_client()
        xai_page = await client.models.list()
        xai_models: list[dict] = []
        for m in xai_page.data:
            mid = m.id
            if not mid.startswith("grok"):
                continue
            xai_models.append({"id": mid, "name": mid})
        result["xai"] = sorted(xai_models, key=lambda x: x["name"])
        log.info("Loaded %d xAI models", len(result["xai"]))
    except Exception as e:
        log.warning("Failed to load xAI models: %s", e)
        result["xai"] = []

    _models_cache["data"] = result
    _models_cache["ts"] = now
    return result

# ── helpers ──────────────────────────────────────────────────────────────────

def extract_pdf_text(pdf_bytes: bytes) -> str:
    """Extract text from PDF bytes using pymupdf."""
    text_parts = []
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"[Page {page_num}]\n{page_text}")
        doc.close()
    except Exception as e:
        log.warning("PDF text extraction failed: %s", e)
        return "[Could not extract text from PDF]"
    return "\n".join(text_parts) if text_parts else "[PDF contained no extractable text]"


def list_projects():
    root = Path(DOCS_ROOT)
    if not root.exists():
        return []
    return sorted([d.name for d in root.iterdir() if d.is_dir()])


def browse_path(rel_path: str = ""):
    """Return immediate subfolders for a given relative path under DOCS_ROOT."""
    root = Path(DOCS_ROOT).resolve()
    target = (root / rel_path).resolve() if rel_path else root
    # path traversal guard
    if not str(target).startswith(str(root)):
        return []
    if not target.exists() or not target.is_dir():
        return []

    children = []
    for d in sorted(target.iterdir()):
        if d.is_dir():
            has_subfolders = any(sub.is_dir() for sub in d.iterdir())
            file_count = sum(1 for f in d.iterdir() if f.is_file())
            children.append({
                "name": d.name,
                "path": str(d.relative_to(root)).replace("\\", "/"),
                "has_children": has_subfolders,
                "file_count": file_count,
            })
    return children


# ── check optional dependencies once at import time ──────────────────────────
try:
    import docx as _docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl as _openpyxl
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

try:
    import xlrd as _xlrd
    HAS_XLS = True
except ImportError:
    HAS_XLS = False


def load_folder(folder_path: Path):
    """Return a list of content blocks (text or image/pdf) for all files in folder."""
    SUPPORTED_TEXT  = {".txt", ".md", ".csv", ".json", ".xml", ".html", ".py",
                       ".js", ".ts", ".java", ".c", ".cpp", ".cs", ".go", ".rs"}
    SUPPORTED_IMAGE = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
    SUPPORTED_DIRECT = {".pdf"}

    blocks = []
    total_size = 0
    files  = sorted(folder_path.rglob("*"))

    for f in files:
        if not f.is_file():
            continue

        # ── safety: skip files that are too large ──
        try:
            fsize = f.stat().st_size
        except OSError:
            continue
        if fsize > MAX_FILE_SIZE:
            log.warning("Skipping %s — exceeds %d MB limit", f.name, MAX_FILE_SIZE // (1024 * 1024))
            continue
        if total_size + fsize > MAX_TOTAL_SIZE:
            log.warning("Stopping file loading — total context exceeds %d MB", MAX_TOTAL_SIZE // (1024 * 1024))
            break
        total_size += fsize

        ext = f.suffix.lower()
        # show path relative to the project folder for clarity
        rel_name = str(f.relative_to(folder_path)).replace("\\", "/")

        # ── plain text ──
        if ext in SUPPORTED_TEXT:
            try:
                content = f.read_text(encoding="utf-8", errors="ignore")
                blocks.append({"type": "text", "text": f"--- FILE: {rel_name} ---\n{content}\n"})
            except Exception as e:
                log.warning("Failed to read text file %s: %s", rel_name, e)

        # ── images ──
        elif ext in SUPPORTED_IMAGE:
            try:
                data    = base64.standard_b64encode(f.read_bytes()).decode()
                mime, _ = mimetypes.guess_type(str(f))
                mime    = mime or "image/jpeg"
                blocks.append({"type": "image", "source": {"type": "base64", "media_type": mime, "data": data}})
                blocks.append({"type": "text", "text": f"(image above: {rel_name})\n"})
            except Exception as e:
                log.warning("Failed to read image %s: %s", rel_name, e)

        # ── PDF ──
        elif ext in SUPPORTED_DIRECT:
            try:
                raw  = f.read_bytes()
                data = base64.standard_b64encode(raw).decode()
                blocks.append({
                    "type": "document",
                    "source": {"type": "base64", "media_type": "application/pdf", "data": data},
                    "_pdf_text": extract_pdf_text(raw),   # pre-extract for models that need it
                    "_filename": rel_name,
                })
                blocks.append({"type": "text", "text": f"(document above: {rel_name})\n"})
            except Exception as e:
                log.warning("Failed to read PDF %s: %s", rel_name, e)

        # ── Word (.docx) — paragraphs + tables ──
        elif ext == ".docx" and HAS_DOCX:
            try:
                import docx
                doc  = docx.Document(str(f))
                parts = []
                for element in doc.element.body:
                    # paragraphs
                    if element.tag.endswith('}p'):
                        for para in doc.paragraphs:
                            if para._element is element:
                                if para.text.strip():
                                    parts.append(para.text)
                                break
                    # tables
                    elif element.tag.endswith('}tbl'):
                        for table in doc.tables:
                            if table._element is element:
                                rows_text = []
                                for row in table.rows:
                                    cells = [cell.text.strip() for cell in row.cells]
                                    rows_text.append("\t".join(cells))
                                parts.append("\n".join(rows_text))
                                break
                blocks.append({"type": "text", "text": f"--- FILE: {rel_name} ---\n" + "\n".join(parts) + "\n"})
            except Exception as e:
                log.warning("Failed to read Word file %s: %s", rel_name, e)

        # ── Excel (.xlsx) ──
        elif ext == ".xlsx" and HAS_XLSX:
            try:
                import openpyxl
                wb   = openpyxl.load_workbook(str(f), data_only=True)
                text = ""
                for sheet in wb.sheetnames:
                    ws   = wb[sheet]
                    text += f"[Sheet: {sheet}]\n"
                    for row in ws.iter_rows(values_only=True):
                        text += "\t".join(str(c) if c is not None else "" for c in row) + "\n"
                blocks.append({"type": "text", "text": f"--- FILE: {rel_name} ---\n{text}\n"})
            except Exception as e:
                log.warning("Failed to read Excel (.xlsx) file %s: %s", rel_name, e)

        # ── Legacy Excel (.xls) ──
        elif ext == ".xls" and HAS_XLS:
            try:
                import xlrd
                wb   = xlrd.open_workbook(str(f))
                text = ""
                for sheet in wb.sheet_names():
                    ws   = wb.sheet_by_name(sheet)
                    text += f"[Sheet: {sheet}]\n"
                    for rx in range(ws.nrows):
                        text += "\t".join(str(c) for c in ws.row_values(rx)) + "\n"
                blocks.append({"type": "text", "text": f"--- FILE: {rel_name} ---\n{text}\n"})
            except Exception as e:
                log.warning("Failed to read Excel (.xls) file %s: %s", rel_name, e)

    log.info("Loaded %d content blocks from %s (%.1f MB)", len(blocks), folder_path.name, total_size / (1024 * 1024))
    return blocks


# ── request / response models ─────────────────────────────────────────────────

class ChatRequest(BaseModel):
    project:       Optional[str] = None
    question:       str
    provider:       str  # "anthropic" | "openai" | "gemini" | "xai"
    model:          str  # actual model ID, e.g. "claude-sonnet-4-5"
    history:        Optional[list] = []
    web_search:     Optional[bool] = True


# ── routes ────────────────────────────────────────────────────────────────────

@app.get("/api/health")
async def health_check():
    """Health check endpoint for monitoring (Uptime Kuma, Docker healthcheck)."""
    return {"status": "ok"}


@app.get("/api/projects")
async def get_projects():
    return {"projects": list_projects()}


@app.get("/api/browse")
async def browse_folder(path: str = ""):
    """Return subfolders for a given relative path, enabling tree navigation."""
    children = browse_path(path)
    return {"path": path, "children": children}


@app.get("/api/models")
async def get_models():
    return await fetch_available_models()


@app.post("/api/chat")
async def chat(req: ChatRequest):
    log.info("Chat request: provider=%s model=%s project=%s history_len=%d",
             req.provider, req.model, req.project, len(req.history or []))

    # ── load project files ────────────────────────────────────────────────────
    context_blocks = []
    if req.project:
        folder = (Path(DOCS_ROOT) / req.project).resolve()
        if not str(folder).startswith(str(Path(DOCS_ROOT).resolve())):
            raise HTTPException(status_code=400, detail="Invalid project path")
        if not folder.exists():
            raise HTTPException(status_code=404, detail="Project folder not found")
        context_blocks = await asyncio.to_thread(load_folder, folder)

    # ── build conversation ────────────────────────────────────────────────────
    # Mirrors how Claude/ChatGPT web works:
    #   • First message  → files + question together as the user message
    #   • Follow-ups     → files injected into the *first* history user message,
    #                       current message is text-only.
    # Result: the AI sees files exactly once in the conversation.
    history = list(req.history or [])

    if history and context_blocks:
        for i, h in enumerate(history):
            if h["role"] == "user":
                history[i] = {
                    "role": "user",
                    "content": context_blocks + [{"type": "text", "text": h["content"]}],
                }
                break
        user_content = [{"type": "text", "text": req.question}]
    else:
        user_content = context_blocks + [{"type": "text", "text": req.question}]

    system_prompt = (
        "You are a helpful assistant with access to the user's project documents. "
        "Answer questions based on the provided documents. "
        "When asked to write something similar to existing documents, match their style, "
        "tone, structure, and language. Be precise and professional."
    )

    # ── helpers: convert content (string or block-array) per provider ──────────
    def _clean_anthropic(content):
        if isinstance(content, list):
            return [{k: v for k, v in b.items() if not k.startswith("_")} for b in content]
        return content

    def _to_oai(block):
        if block["type"] == "text":
            return {"type": "text", "text": block["text"]}
        if block["type"] == "image":
            src = block["source"]
            return {"type": "image_url", "image_url": {"url": f"data:{src['media_type']};base64,{src['data']}"}}
        if block["type"] == "document":
            return {"type": "text", "text": f"--- FILE: {block.get('_filename','document.pdf')} ---\n{block.get('_pdf_text','[PDF]')}\n"}
        return {"type": "text", "text": ""}

    def _to_oai_resp(block):
        if block["type"] == "text":
            return {"type": "input_text", "text": block["text"]}
        if block["type"] == "image":
            src = block["source"]
            return {"type": "input_image", "image_url": f"data:{src['media_type']};base64,{src['data']}"}
        if block["type"] == "document":
            return {"type": "input_text", "text": f"--- FILE: {block.get('_filename','document.pdf')} ---\n{block.get('_pdf_text','[PDF]')}\n"}
        return {"type": "input_text", "text": ""}

    def _oai_content(content, converter):
        if isinstance(content, list):
            return [converter(b) for b in content]
        return content

    def _gemini_parts(content):
        if isinstance(content, str):
            return [gtypes.Part.from_text(text=content)]
        parts = []
        for block in content:
            if block["type"] == "text":
                parts.append(gtypes.Part.from_text(text=block["text"]))
            elif block["type"] == "image":
                src = block["source"]
                parts.append(gtypes.Part.from_bytes(data=base64.b64decode(src["data"]), mime_type=src["media_type"]))
            elif block["type"] == "document":
                parts.append(gtypes.Part.from_bytes(data=base64.b64decode(block["source"]["data"]), mime_type="application/pdf"))
        return parts

    try:
        # ── Claude ───────────────────────────────────────────────────────────
        if req.provider == "anthropic":
            client = get_anthropic_client()
            messages = []
            for h in history:
                messages.append({"role": h["role"], "content": _clean_anthropic(h["content"])})
            messages.append({"role": "user", "content": _clean_anthropic(user_content)})
            kwargs = dict(model=req.model, max_tokens=16000, system=system_prompt, messages=messages)
            if req.web_search:
                kwargs["tools"] = [{"type": "web_search_20250305", "name": "web_search", "max_uses": 5}]
            response = await client.messages.create(**kwargs)
            text_parts = [b.text for b in response.content if hasattr(b, "text")]
            return {"answer": "\n\n".join(text_parts) if text_parts else "No response generated."}

        # ── GPT-4o ───────────────────────────────────────────────────────────
        elif req.provider == "openai":
            client = get_openai_client()
            if req.web_search:
                messages = []
                for h in history:
                    messages.append({"role": h["role"], "content": _oai_content(h["content"], _to_oai_resp)})
                messages.append({"role": "user", "content": [_to_oai_resp(b) for b in user_content]})
                response = await client.responses.create(
                    model=req.model, instructions=system_prompt, input=messages,
                    tools=[{"type": "web_search_preview"}],
                )
                return {"answer": response.output_text}

            messages = [{"role": "system", "content": system_prompt}]
            for h in history:
                messages.append({"role": h["role"], "content": _oai_content(h["content"], _to_oai)})
            messages.append({"role": "user", "content": [_to_oai(b) for b in user_content]})
            response = await client.chat.completions.create(model=req.model, messages=messages, max_tokens=16000)
            return {"answer": response.choices[0].message.content}

        # ── Gemini (sync SDK — run in thread) ────────────────────────────────
        elif req.provider == "gemini":
            client = get_gemini_client()
            config = gtypes.GenerateContentConfig(
                system_instruction=system_prompt, max_output_tokens=16000,
                tools=[gtypes.Tool(google_search=gtypes.GoogleSearch())] if req.web_search else None,
            )
            gemini_history: list = []
            for h in history:
                role = "user" if h["role"] == "user" else "model"
                gemini_history.append(gtypes.Content(role=role, parts=_gemini_parts(h["content"])))
            current_parts = _gemini_parts(user_content)

            def _gemini_call():
                if gemini_history:
                    session = client.chats.create(model=req.model, config=config, history=gemini_history)
                    return session.send_message(current_parts)
                return client.models.generate_content(model=req.model, contents=current_parts, config=config)

            response = await asyncio.to_thread(_gemini_call)
            return {"answer": response.text}

        # ── xAI (Grok) ──────────────────────────────────────────────────────
        elif req.provider == "xai":
            client = get_xai_client()
            messages = [{"role": "system", "content": system_prompt}]
            for h in history:
                messages.append({"role": h["role"], "content": _oai_content(h["content"], _to_oai)})
            messages.append({"role": "user", "content": [_to_oai(b) for b in user_content]})
            kwargs = dict(model=req.model, messages=messages, max_tokens=16000)
            if req.web_search:
                kwargs["extra_body"] = {"search_parameters": {"mode": "auto", "max_search_results": 5}}
            response = await client.chat.completions.create(**kwargs)
            return {"answer": response.choices[0].message.content}

        else:
            raise HTTPException(status_code=400, detail="Unknown provider")

    except HTTPException:
        raise
    except Exception as e:
        log.error("AI request failed: provider=%s model=%s error=%s", req.provider, req.model, e, exc_info=True)
        raise HTTPException(status_code=502, detail=f"AI provider error: {str(e)}")


# serve frontend
frontend_path = Path(__file__).parent.parent / "frontend"
if frontend_path.exists():
    app.mount("/", StaticFiles(directory=str(frontend_path), html=True), name="static")
